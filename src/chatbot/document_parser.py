import os
import logging
from typing import List, Dict, Any
from pathlib import Path

from pypdf import PdfReader
import docx
from unstructured.partition.auto import partition
from unstructured.chunking.title import chunk_by_title

logger = logging.getLogger(__name__)

class DocumentParser:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
    def parse_document(self, file_path: str) -> List[Dict[str, Any]]:
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif file_extension == '.txt':
                return self._parse_txt(file_path)
            else:
                return self._parse_with_unstructured(file_path)
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        chunks = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            full_text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                full_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
            text_chunks = self._chunk_text(full_text)
            
            for i, chunk in enumerate(text_chunks):
                chunks.append({
                    'content': chunk,
                    'metadata': {
                        'source': str(file_path),
                        'chunk_id': i,
                        'file_type': 'pdf',
                        'total_pages': len(pdf_reader.pages)
                    }
                })
        
        return chunks
    
    def _parse_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        chunks = []
        
        doc = docx.Document(file_path)
        full_text = ""
        
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        text_chunks = self._chunk_text(full_text)
        
        for i, chunk in enumerate(text_chunks):
            chunks.append({
                'content': chunk,
                'metadata': {
                    'source': str(file_path),
                    'chunk_id': i,
                    'file_type': 'docx',
                    'total_paragraphs': len(doc.paragraphs)
                }
            })
        
        return chunks
    
    def _parse_txt(self, file_path: Path) -> List[Dict[str, Any]]:
        chunks = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            full_text = file.read()
        
        text_chunks = self._chunk_text(full_text)
        
        for i, chunk in enumerate(text_chunks):
            chunks.append({
                'content': chunk,
                'metadata': {
                    'source': str(file_path),
                    'chunk_id': i,
                    'file_type': 'txt'
                }
            })
        
        return chunks
    
    def _parse_with_unstructured(self, file_path: Path) -> List[Dict[str, Any]]:
        chunks = []
        
        elements = partition(filename=str(file_path))
        chunked_elements = chunk_by_title(elements, max_characters=self.chunk_size)
        
        for i, element in enumerate(chunked_elements):
            chunks.append({
                'content': element.text,
                'metadata': {
                    'source': str(file_path),
                    'chunk_id': i,
                    'file_type': file_path.suffix.lower().replace('.', ''),
                    'element_type': element.category if hasattr(element, 'category') else 'unknown'
                }
            })
        
        return chunks
    
    def _chunk_text(self, text: str) -> List[str]:
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk = ' '.join(chunk_words)
            chunks.append(chunk)
            
        return chunks
    
    def parse_directory(self, directory_path: str, file_extensions: List[str] = None) -> List[Dict[str, Any]]:
        if file_extensions is None:
            file_extensions = ['.pdf', '.docx', '.doc', '.txt', '.md', '.html', '.csv']
        
        directory_path = Path(directory_path)
        all_chunks = []
        
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_extensions:
                try:
                    chunks = self.parse_document(str(file_path))
                    all_chunks.extend(chunks)
                    logger.info(f"Parsed {len(chunks)} chunks from {file_path}")
                except Exception as e:
                    logger.error(f"Failed to parse {file_path}: {str(e)}")
        
        return all_chunks